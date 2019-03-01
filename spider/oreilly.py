# coding:utf-8
'''
Created on 2015-12-28
@author: Jay
target:
this is my first try to write a long programme
'''
import urllib
import urllib2
import re
import docx


class OreillyPc:
    def __init__(self):
        self.url = 'http://www.oreilly.com.cn/index.php?func=completelist'
        self.url0 = 'http://www.oreilly.com.cn/index.php?func=book&isbn='
        self.url1 = 'http://www.oreilly.com.cn/images/bookcover/'

    def getResponse(self):
        request = urllib2.Request(self.url)
        response = urllib2.urlopen(request)
        return response.read()

    def getLink(self, response):
        pattern = re.compile('<a href="index\.php\?func=book&amp;isbn=(.+?)"')
        temp = re.findall(pattern, response)
        links = [self.url0 + (x if len(x) > 10 else "") for x in temp]
        return links

    def getBook(self, response):
        pattern = re.compile('</div>">(.+?)</a>')
        books = re.findall(pattern, response)
        return books

    def getDate(self, response):
        pattern = re.compile('<td valign="top" align="center" nowrap="nowrap">\n(.+?)</td>')
        temp = re.findall(pattern, response)
        dates = [temp[i] for i in range(len(temp)) if i % 2 == 0]
        return dates

    def getPrice(self, response):
        pattern = re.compile('<td valign="top" align="center" nowrap="nowrap">\n(.+?)</td>')
        temp = re.findall(pattern, response)
        prices = [temp[i] for i in range(len(temp)) if i % 2 == 1]
        return prices

    def comBine(self):
        res = self.getResponse()
        books = self.getBook(res)
        dates = self.getDate(res)
        prices = self.getPrice(res)
        links = self.getLink(res)
        com = map(lambda x, y, z, w: x + '|' + y + '|' + z + '|' + w, books, dates, prices, links)
        return com

    def getImage(self, response, num):
        request = urllib2.Request(response)
        res = urllib2.urlopen(request).read()
        pattern = re.compile(r'images/bookcover/(.+?\.gif)|images/bookcover/(.+?\.jpg)')
        temp = set(re.findall(pattern, res))
        imgurls = [x for t in temp for x in t if len(x) > 0]
        if len(imgurls) > 0:
            # only get large pic
            if len(imgurls) == 1:
                i = 0
            else:
                i = 1 if len(imgurls[0]) < len(imgurls[1]) else 0
            imgurl = self.url1 + imgurls[i]
            name = str(num) + '_' + str(i) + '_' + imgurls[i]
            urllib.urlretrieve(imgurl, '/Volumes/File/oreilly/%s' % name)
        else:
            name = 'NO_PIC'
        return name

    def getCatalog(self, response):
        request = urllib2.Request(response)
        res = urllib2.urlopen(request).read()
        pattern = re.compile(r'<ol>\n((.+\n)+?)</ol>')
        temp = re.search(pattern, res)
        if temp is None:
            log = 'NO_Catalog'
        else:
            log = temp.group(1).replace('</li>', '').replace('<li>', '')
        return log

    def getBookinfo(self, response):
        request = urllib2.Request(response)
        temp0 = urllib2.urlopen(request).read()
        temp1 = re.search(r'products_details((\S+\s+)+?)tab2', temp0)
        if temp1:
            res = temp1.group(0)
            pattern = re.compile(r'<div>(.+?)</div>')
            temp2 = [re.sub(r'<.+?>', '', i) for i in re.findall(pattern, res)]
            temp3 = [re.sub(r'&nbsp;', '', i) for i in temp2]
            info = [i for i in temp3 if len(i) > 0]
        else:
            info = 'NO_INFO'
        return info

    def getSummary(self, response):
        request = urllib2.Request(response)
        res = urllib2.urlopen(request).read()
        pattern = re.compile(r"<div id='detail-description-container'>(.+?)</div>")
        temp = re.search(pattern, re.sub(r'\n', '', res))
        if temp:
            summary = re.sub('\n+', '', temp.group(1).replace('<br />', ''))
        else:
            summary = 'NO_SUMMARY'
        return summary


doc = Document()
doc.add_heading("O'Reilly Books", level=0)

mypc = OreillyPc()
myres = mypc.getResponse()
mylinks = mypc.getLink(myres)
mybooks = mypc.getBook(myres)
mydates = mypc.getDate(myres)
myprice = mypc.getPrice(myres)

for i in range(len(mybooks)):
    book = mybooks[i]
    doc.add_heading(book.decode('utf-8'), level=1)
    doc.add_paragraph('')

    name = mypc.getImage(mylinks[i], i)
    try:
        if name == 'NO_PIC':
            doc.add_paragraph('NO_PIC')
            print 'NO_PIC:' + str(i)
        else:
            p = doc.add_picture('/Volumes/File/oreilly/' + name, width=Inches(2), height=Inches(2.6))
    except Exception, e:
        print 'pic_problem:' + str(i)

    infos = mypc.getBookinfo(mylinks[i])
    info_str = ''
    for j in infos:
        info_str = info_str + j + '\n'
    try:
        doc.add_paragraph(info_str.decode('utf-8'))
    except Exception, e:
        print 'infos_problem:' + str(i)

    summary = mypc.getSummary(mylinks[i])
    try:
        doc.add_paragraph(summary.decode('utf-8'))
    except Exception, e:
        print 'summary_problem:' + str(i)
    doc.add_page_break()
doc.save('./oreilly5.0.docx')

